from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Comprehensive Backend Architecture Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Header Info
doc.add_paragraph('Project: Liquid AI LLM Chatbot Assignment').bold = True
doc.add_paragraph('Course: SE803 Software Project Management').bold = True

doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    "The backend of the chatbot system is designed to be highly scalable, modular, and secure. "
    "It is built using the FastAPI framework in Python and acts as the central hub connecting the "
    "frontend interface, the persistent relational database, the semantic vector database, and the "
    "local Large Language Model (LLM)."
)

doc.add_heading('2. Core Components and Modules', level=1)

# FastAPI & Routing
doc.add_heading('2.1 FastAPI Framework & Routing', level=2)
doc.add_paragraph(
    "FastAPI is utilized for high-performance, asynchronous API routing. The routing is modularized "
    "into dedicated files (e.g., routes_auth.py, routes_threads.py, routes_health.py). This separation "
    "of concerns ensures clean API endpoint management."
)

# Authentication
doc.add_heading('2.2 Authentication & Security', level=2)
doc.add_paragraph(
    "The authentication system (auth.py) is implemented using JSON Web Tokens (JWT) and bcrypt password hashing. "
    "Users must register and log in to receive a JWT Bearer token. This token is subsequently required "
    "in the Authorization header for all API requests, ensuring that chatbot interactions and history "
    "retrieval are strictly tied to the authenticated user's session."
)

# Database
doc.add_heading('2.3 Relational Database Persistence (SQLite/SQLAlchemy)', level=2)
doc.add_paragraph(
    "Persistent storage for user credentials and exact chat transcripts is handled by SQLite and "
    "managed through SQLAlchemy ORM (models.py, database.py). The schema dictates relationships "
    "where a User owns multiple Threads, and each Thread owns multiple Message Items. This structural "
    "integrity guarantees that a returning user can seamlessly fetch their entire previous chat history."
)

# ChatKit
doc.add_heading('2.4 ChatKit Protocol Handling', level=2)
doc.add_paragraph(
    "The backend implements a custom OpenAI ChatKit server endpoint (chatkit_server.py, chatkit_store.py). "
    "Instead of routing messages directly to OpenAI, the server intercepts the standard ChatKit protocol payloads, "
    "extracts the prompt, logs the interaction in the SQLite database, and hands the input off to the custom LLM pipeline."
)

# LLM integration
doc.add_heading('2.5 Language Model Integration (LangChain & Ollama)', level=2)
doc.add_paragraph(
    "The core conversational logic is powered by LangChain (llm.py). The backend uses the ChatOllama client "
    "to query the open-source Liquid AI LFM model (hf.co/LiquidAI/LFM2.5-1.2B-Instruct-GGUF) running locally. "
    "LangChain orchestrates the prompt construction, seamlessly injecting retrieved context before generating the response."
)

# Qdrant Memory
doc.add_heading('2.6 Semantic Memory (Qdrant)', level=2)
doc.add_paragraph(
    "To provide the chatbot with long-term semantic memory, the backend integrates with Qdrant (qdrant_memory.py). "
    "Every turn in the conversation is embedded into a vector representation and stored in the Qdrant instance. "
    "When a user asks a new question, the system queries Qdrant for contextually relevant past snippets, "
    "enhancing the prompt to the Liquid AI model with contextual awareness beyond the immediate conversation window."
)

doc.add_heading('3. Workflow Summary', level=1)
doc.add_paragraph(
    "1. The user authenticates via FastAPI to receive a JWT.\n"
    "2. The user submits a message via the ChatKit endpoint, validated by the JWT.\n"
    "3. The message is stored in SQLite (chatkit_store.py).\n"
    "4. Semantic context is retrieved from Qdrant (qdrant_memory.py).\n"
    "5. The context and prompt are passed to the Liquid AI model via LangChain (llm.py).\n"
    "6. The response is streamed back to the frontend, simultaneously being logged in both SQLite and Qdrant."
)

doc.save('BACKEND_REPORT.docx')
print("Backend report generated successfully at BACKEND_REPORT.docx")
